"""体检报告文档文本提取服务：支持 .docx / .pdf"""
import io
from typing import List


SUPPORTED_EXTS = {"docx", "pdf"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def extract_text(filename: str, data: bytes) -> str:
    """按扩展名提取文档文本，失败抛出 ValueError"""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"不支持的文件格式：{ext or '未知'}，仅支持 .docx / .pdf")
    if not data:
        raise ValueError("文件内容为空")
    if len(data) > MAX_FILE_SIZE:
        raise ValueError("文件不能超过 10MB")
    try:
        text = _extract_docx(data) if ext == "docx" else _extract_pdf(data)
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"文件解析失败：{e}") from e
    return text


def _extract_docx(data: bytes) -> str:
    """提取 .docx 文本（含段落和表格）"""
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_pdf(data: bytes) -> str:
    """提取 PDF 文本（要求 PDF 自带文字层，扫描件需先 OCR）"""
    import pdfplumber

    parts: List[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()
